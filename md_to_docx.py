"""
md_to_docx.py — shared Markdown to Word converter for PRD Generator

Usage from any initiative folder:
    from md_to_docx import convert
    convert("path/to/source.md", "path/to/output.docx")

Or run directly:
    python3 md_to_docx.py source.md output.docx

Styling standards applied every time:
- Arial font throughout (w:rFonts ascii, hAnsi, cs on all styles and runs)
- All text black (w:color val="000000"; no theme color)
- Line height 1.15
- All headings bold
- No horizontal rules
- Zero bookmarks (strips w:bookmarkStart and w:bookmarkEnd)
"""

import re
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


LINE_SPACING = Pt(11 * 1.15)


# ---------------------------------------------------------------------------
# Font, color, and spacing helpers
# ---------------------------------------------------------------------------

def set_arial_on_run(run):
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    for el in rPr.findall(qn("w:color")):
        rPr.remove(el)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "000000")
    rPr.append(color_el)


def set_arial_on_style(style):
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Arial")
    rFonts.set(qn("w:hAnsi"), "Arial")
    rFonts.set(qn("w:cs"), "Arial")
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    for el in rPr.findall(qn("w:color")):
        rPr.remove(el)
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "000000")
    rPr.append(color_el)


def set_line_spacing(para):
    para.paragraph_format.line_spacing = LINE_SPACING


def apply_styles_to_document(doc):
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3",
                        "Heading 4", "List Paragraph", "Table Grid"]:
        try:
            set_arial_on_style(doc.styles[style_name])
        except KeyError:
            pass
    doc.styles["Normal"].font.size = Pt(11)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].paragraph_format.line_spacing = LINE_SPACING


# ---------------------------------------------------------------------------
# Bookmark stripping
# ---------------------------------------------------------------------------

def strip_bookmarks(doc):
    body = doc.element.body
    for tag in (qn("w:bookmarkStart"), qn("w:bookmarkEnd")):
        for el in list(body.iter(tag)):
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)


# ---------------------------------------------------------------------------
# Inline bold/italic parsing
# ---------------------------------------------------------------------------

def add_inline_runs(paragraph, text):
    pattern = re.compile(r'(\*\*(.+?)\*\*|\*(.+?)\*)')
    last = 0
    for m in pattern.finditer(text):
        start, end = m.start(), m.end()
        if start > last:
            run = paragraph.add_run(text[last:start])
            set_arial_on_run(run)
        if m.group(0).startswith("**"):
            run = paragraph.add_run(m.group(2))
            run.bold = True
            set_arial_on_run(run)
        else:
            run = paragraph.add_run(m.group(3))
            run.italic = True
            set_arial_on_run(run)
        last = end
    if last < len(text):
        run = paragraph.add_run(text[last:])
        set_arial_on_run(run)


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------

def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_separator_row(line):
    return is_table_row(line) and re.match(r'^[\|\s\-:]+$', line.strip())


def parse_table_cells(line):
    stripped = line.strip().strip("|")
    return [cell.strip() for cell in stripped.split("|")]


def build_table(doc, rows):
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row_cells in enumerate(rows):
        for c_idx, cell_text in enumerate(row_cells):
            if c_idx >= col_count:
                break
            cell = table.cell(r_idx, c_idx)
            para = cell.paragraphs[0]
            para.clear()
            add_inline_runs(para, cell_text)
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True
                    set_arial_on_run(run)
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_arial_on_run(run)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def convert(source_path, output_path):
    with open(source_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()
    apply_styles_to_document(doc)

    table_buffer = []
    i = 0

    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip horizontal rules
        if re.match(r'^---+$', line.strip()):
            i += 1
            continue

        # Table detection
        if is_table_row(line):
            table_buffer.append(line)
            i += 1
            continue
        else:
            if table_buffer:
                raw_rows = [r for r in table_buffer if not is_separator_row(r)]
                build_table(doc, [parse_table_cells(r) for r in raw_rows])
                table_buffer = []

        # Headings
        h1 = re.match(r'^# (.+)$', line)
        h2 = re.match(r'^## (.+)$', line)
        h3 = re.match(r'^### (.+)$', line)

        if h1:
            p = doc.add_heading(level=1)
            p.clear()
            run = p.add_run(h1.group(1))
            run.bold = True
            set_arial_on_run(run)
            set_line_spacing(p)
            i += 1
            continue

        if h2:
            p = doc.add_heading(level=2)
            p.clear()
            run = p.add_run(h2.group(1))
            run.bold = True
            set_arial_on_run(run)
            set_line_spacing(p)
            i += 1
            continue

        if h3:
            p = doc.add_heading(level=3)
            p.clear()
            run = p.add_run(h3.group(1))
            run.bold = True
            set_arial_on_run(run)
            set_line_spacing(p)
            i += 1
            continue

        # Bullet list
        bullet = re.match(r'^[-*]\s+(.+)$', line)
        if bullet:
            p = doc.add_paragraph(style="List Paragraph")
            pPr = p._p.get_or_add_pPr()
            numPr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "1")
            numPr.append(ilvl)
            numPr.append(numId)
            pPr.append(numPr)
            add_inline_runs(p, bullet.group(1))
            set_line_spacing(p)
            i += 1
            continue

        # Numbered list
        numbered = re.match(r'^\d+\.\s+(.+)$', line)
        if numbered:
            p = doc.add_paragraph(style="List Paragraph")
            add_inline_runs(p, numbered.group(1))
            set_line_spacing(p)
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Default: normal paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, line)
        set_line_spacing(p)
        i += 1

    # Flush any remaining table
    if table_buffer:
        raw_rows = [r for r in table_buffer if not is_separator_row(r)]
        build_table(doc, [parse_table_cells(r) for r in raw_rows])

    strip_bookmarks(doc)
    doc.save(output_path)
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python3 md_to_docx.py source.md output.docx")
        sys.exit(1)
    convert(sys.argv[1], sys.argv[2])
